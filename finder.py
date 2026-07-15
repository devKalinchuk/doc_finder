#!/usr/bin/env python3
import argparse
import re
from pathlib import Path
from docx import Document


TARGET_FOLDER = "."  # Папка за замовчуванням

# ANSI-коди для підсвічування тексту червоним кольором у терміналі
RED = "\033[91m"
RESET = "\033[0m"


def extract_paragraphs(doc):
    """
    Генератор, який повертає всі абзаци документа (Paragraph-об'єкти):
    основний текст, текст у таблицях та колонтитулах.
    """
    for para in doc.paragraphs:
        yield para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs

    for section in doc.sections:
        if section.header:
            yield from section.header.paragraphs
        if section.footer:
            yield from section.footer.paragraphs


def build_pattern(terms, case_sensitive):
    """
    Створює один регулярний вираз-«або» для всіх слів/фраз пошуку.
    Кожен термін екранується (re.escape), тому спецсимволи сприймаються
    буквально, а фрази з лапок (де слова вже з'єднані пробілом) шукаються
    як єдиний нерозривний рядок.
    """
    escaped_terms = [re.escape(term) for term in terms]
    pattern_str = "|".join(escaped_terms)
    flags = 0 if case_sensitive else re.IGNORECASE
    return re.compile(pattern_str, flags)


def highlight(text, pattern):
    """Підсвічує всі знайдені збіги червоним кольором для виводу в термінал."""
    return pattern.sub(lambda m: f"{RED}{m.group(0)}{RESET}", text)


def find_text_in_folder(folder_path_str, terms, case_sensitive=False, recursive=False):
    folder = Path(folder_path_str)

    if not folder.exists() or not folder.is_dir():
        print(f"Помилка: Шлях '{folder}' не існує або не є папкою.")
        return

    glob_method = folder.rglob if recursive else folder.glob
    docx_files = [
        f for f in glob_method("*.docx")
        if not f.name.startswith("~$")
    ]

    if not docx_files:
        print("У вказаній папці не знайдено документів .docx")
        return

    pattern = build_pattern(terms, case_sensitive)
    total_matches = 0

    for file_path in docx_files:
        try:
            doc = Document(file_path)
            file_header_printed = False

            for para in extract_paragraphs(doc):
                text = para.text
                if text.strip() and pattern.search(text):
                    if not file_header_printed:
                        print(f"📄 Файл: {file_path.name}")
                        file_header_printed = True
                    print(f"   {highlight(text, pattern)}")
                    total_matches += 1

            if file_header_printed:
                print("-" * 40)

        except Exception as e:
            print(f"⚠️ Не вдалося прочитати файл {file_path.name}: {e}")

    if total_matches == 0:
        print("Збігів не знайдено.")


def parse_args():
    parser = argparse.ArgumentParser(
        prog="finder",
        description="Пошук тексту в документах .docx з підсвічуванням збігів червоним кольором."
    )
    parser.add_argument(
        "query",
        nargs="+",
        help='Текст для пошуку. Кілька слів шукаються кожне окремо (finder слово1 слово2), '
             'а текст "у лапках" шукається як одна нерозривна фраза (finder "слово1 слово2").'
    )
    parser.add_argument(
        "-c", "--case-sensitive",
        action="store_true",
        help="Точний пошук з урахуванням регістру (за замовчуванням регістр ігнорується)."
    )
    parser.add_argument(
        "-f", "--folder",
        default=TARGET_FOLDER,
        help="Папка з документами .docx (за замовчуванням поточна папка).",
    )
    parser.add_argument(
        "-r", "--recursive",
        action="store_true",
        help="Шукати документи також у підпапках.",
    )
    return parser.parse_args()


# --- Запуск ---
if __name__ == "__main__":
    args = parse_args()
    print(
        f"Пошук: {' | '.join(args.query)}\n"
        f"Папка: '{Path(args.folder).resolve()}'"
        f"{' (рекурсивно)' if args.recursive else ''}"
        f"{' [точний регістр]' if args.case_sensitive else ''}\n"
    )
    find_text_in_folder(
        args.folder,
        args.query,
        case_sensitive=args.case_sensitive,
        recursive=args.recursive,
    )