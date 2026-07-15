from pathlib import Path
import re
from docx import Document


TARGET_FOLDER = "."  # Можна змінити на будь-який інший шлях до папки з документами .docx

# Регулярний вираз для пошуку українських номерів
PHONE_PATTERN = re.compile(
    r'(?:\+?38)?\s*\(?0\d{2}\)?[-.\s]*\d{3}[-.\s]*\d{2}[-.\s]*\d{2}\b'
)


def extract_text_sources(doc):
    """
    Генератор, який послідовно повертає весь текст із різних частин документа:
    основні абзаци, таблиці та колонтитули.
    """
    # 1. Текст з основних абзаців
    for para in doc.paragraphs:
        yield para.text

    # 2. Текст із таблиць
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.text

    # 3. Текст із колонтитулів (якщо вони існують)
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                yield para.text
        if section.footer:
            for para in section.footer.paragraphs:
                yield para.text


def find_phones_in_folder(folder_path_str):
    folder = Path(folder_path_str)

    if not folder.exists() or not folder.is_dir():
        print(f"Помилка: Шлях '{folder}' не існує або не є папкою.")
        return

    # Шукаємо всі файли .docx за допомогою rglob (рекурсивний пошук) або glob (тільки в цій папці)
    # Використовуємо glob, щоб шукати лише в поточній папці (без підпапок)
    docx_files = [
        f for f in folder.glob("*.docx")
        if not f.name.startswith("~$")
    ]

    if not docx_files:
        print("У вказаній папці не знайдено документів .docx")
        return

    for file_path in docx_files:
        try:
            doc = Document(file_path)
            phones_in_file = []

            # Збираємо текст з усіх джерел у документі
            for text in extract_text_sources(doc):
                if text.strip():
                    matches = PHONE_PATTERN.findall(text)
                    for match in matches:
                        cleaned_phone = match.strip()
                        if cleaned_phone not in phones_in_file:
                            phones_in_file.append(cleaned_phone)

            # Вивід результатів (лише якщо знайдено телефони)
            if phones_in_file:
                # file_path.name повертає тільки ім'я файлу (наприклад, "document.docx")
                print(f"📄 Файл: {file_path.name}")
                for phone in phones_in_file:
                    print(f"   📞 {phone}")
                print("-" * 40)

        except Exception as e:
            print(f"⚠️ Не вдалося прочитати файл {file_path.name}: {e}")


# --- Запуск ---
if __name__ == "__main__":
    print(f"Пошук телефонів у папці '{Path(TARGET_FOLDER).resolve()}'...\n")
    find_phones_in_folder(TARGET_FOLDER)